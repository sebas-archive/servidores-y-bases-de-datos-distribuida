"""Generador del documento de entrega de la Fase 2 (.docx, normas APA 7).

Genera docs/FASE-02/Informe-Fase-2-Sistemas-Distribuidos.docx combinando:
portada, declaración de uso de herramientas de IA, introducción,
requerimientos, resumen del sistema, instalación y configuración,
evidencias (con etiquetas [IMAGEN: ...] para que el equipo inserte las
capturas), conclusiones y referencias bibliográficas.

Requisitos: pip install python-docx
Uso: python3 tools/generar-informe-fase2.py
"""

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

SALIDA = "docs/FASE-02/Informe-Fase-2-Sistemas-Distribuidos.docx"

# ---------------------------------------------------------------------------
# Estilo base (APA 7: Times New Roman 12, interlineado doble, márgenes 2,54)
# ---------------------------------------------------------------------------

doc = Document()

section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 2.0
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(0)

for nivel, alineacion, cursiva in (("Title", WD_ALIGN_PARAGRAPH.CENTER, False),
                                   ("Heading 1", WD_ALIGN_PARAGRAPH.CENTER, False),
                                   ("Heading 2", WD_ALIGN_PARAGRAPH.LEFT, False),
                                   ("Heading 3", WD_ALIGN_PARAGRAPH.LEFT, True)):
    estilo = doc.styles[nivel]
    estilo.font.name = "Times New Roman"
    estilo.font.size = Pt(12)
    estilo.font.bold = True
    estilo.font.italic = cursiva
    estilo.font.color.rgb = None  # negro (por defecto en Heading es azul)
    estilo.paragraph_format.alignment = alineacion
    estilo.paragraph_format.line_spacing = 2.0
    estilo.paragraph_format.space_before = Pt(0)
    estilo.paragraph_format.space_after = Pt(0)
    estilo.paragraph_format.keep_with_next = True

doc.styles["Title"].font.size = Pt(14)


# Número de página en el encabezado (esquina superior derecha)
encabezado = section.header
parrafo_numero = encabezado.paragraphs[0]
parrafo_numero.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = parrafo_numero.add_run()
fld_ini = OxmlElement("w:fldChar")
fld_ini.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText")
instr.set(qn("xml:space"), "preserve")
instr.text = "PAGE"
fld_fin = OxmlElement("w:fldChar")
fld_fin.set(qn("w:fldCharType"), "end")
run._r.append(fld_ini)
run._r.append(instr)
run._r.append(fld_fin)
run.font.name = "Times New Roman"
run.font.size = Pt(12)


# ---------------------------------------------------------------------------
# Ayudantes de contenido
# ---------------------------------------------------------------------------

def h1(texto):
    doc.add_heading(texto, level=1)


def h2(texto):
    doc.add_heading(texto, level=2)


def p(texto, sangria=True):
    """Párrafo de cuerpo APA (sangría de primera línea 1,27 cm)."""
    par = doc.add_paragraph(texto)
    if sangria:
        par.paragraph_format.first_line_indent = Cm(1.27)
    return par


def item(texto):
    """Elemento de lista (RF/RNF) con sangría francesa."""
    par = doc.add_paragraph(texto)
    par.paragraph_format.left_indent = Cm(1.27)
    par.paragraph_format.first_line_indent = Cm(-1.27)
    return par


def codigo(lineas):
    """Bloque de comandos en fuente monoespaciada, a espacio sencillo."""
    for linea in lineas:
        par = doc.add_paragraph(linea)
        par.paragraph_format.left_indent = Cm(1.27)
        par.paragraph_format.line_spacing = 1.0
        par.paragraph_format.space_after = Pt(0)
        for run in par.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(10)


def etiqueta_imagen(texto):
    """Recuadro marcador que el equipo sustituye por la captura real."""
    tabla = doc.add_table(rows=1, cols=1)
    tabla.style = "Table Grid"
    celda = tabla.cell(0, 0)
    celda.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    par = celda.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.line_spacing = 1.0
    run = par.add_run(texto)
    run.italic = True
    fila = tabla.rows[0]
    fila.height = Cm(2.5)
    fila.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def figura(numero, titulo, etiqueta, nota=None):
    """Figura APA 7: número (negrita) + título (cursiva), recuadro y nota."""
    par = doc.add_paragraph()
    run = par.add_run(f"Figura {numero}")
    run.bold = True
    par.paragraph_format.keep_with_next = True
    par2 = doc.add_paragraph()
    run2 = par2.add_run(titulo)
    run2.italic = True
    par2.paragraph_format.keep_with_next = True
    etiqueta_imagen(etiqueta)
    if nota:
        par3 = doc.add_paragraph()
        run3 = par3.add_run(f"Nota. {nota}")
        run3.font.size = Pt(11)
        par3.paragraph_format.line_spacing = 1.0


def titulo_tabla(numero, titulo):
    """Encabezado de tabla APA 7: número (negrita) + título (cursiva)."""
    par = doc.add_paragraph()
    run = par.add_run(f"Tabla {numero}")
    run.bold = True
    par.paragraph_format.keep_with_next = True
    par2 = doc.add_paragraph()
    run2 = par2.add_run(titulo)
    run2.italic = True
    par2.paragraph_format.keep_with_next = True


def tabla(encabezados, filas, anchos=None, nota=None):
    """Tabla sencilla con bordes; encabezado en negrita; nota al pie."""
    t = doc.add_table(rows=1 + len(filas), cols=len(encabezados))
    t.style = "Table Grid"
    for j, texto in enumerate(encabezados):
        celda = t.cell(0, j)
        celda.paragraphs[0].paragraph_format.line_spacing = 1.0
        run = celda.paragraphs[0].add_run(texto)
        run.bold = True
        if anchos:
            celda.width = Cm(anchos[j])
    for i, fila in enumerate(filas, start=1):
        for j, texto in enumerate(fila):
            celda = t.cell(i, j)
            celda.paragraphs[0].paragraph_format.line_spacing = 1.0
            celda.paragraphs[0].add_run(texto)
            if anchos:
                celda.width = Cm(anchos[j])
    if nota:
        par = doc.add_paragraph()
        run = par.add_run(f"Nota. {nota}")
        run.font.size = Pt(11)
        par.paragraph_format.line_spacing = 1.0


def referencia(texto):
    """Referencia APA 7 con sangría francesa."""
    par = doc.add_paragraph(texto)
    par.paragraph_format.left_indent = Cm(1.27)
    par.paragraph_format.first_line_indent = Cm(-1.27)
    return par


# ---------------------------------------------------------------------------
# Portada
# ---------------------------------------------------------------------------

for _ in range(4):
    doc.add_paragraph()

p("Sistema distribuido de registro y consulta de información\n"
  "con replicación y tolerancia a fallos", sangria=False).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

p("Fase 2: implementación, pruebas y evidencias", sangria=False).alignment = WD_ALIGN_PARAGRAPH.CENTER

for _ in range(4):
    doc.add_paragraph()

for linea in ("Jose Camilo Pérez Daza",
              "Sebastián Fernando Revelo Meneses",
              "Tomas Alejandro Santiago Reyes"):
    p(linea, sangria=False).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

for linea in ("Ingeniería de Software",
              "Universidad Manuela Beltrán",
              "Sistemas Distribuidos",
              "Prof. Juan José Osorio Tabares"):
    p(linea, sangria=False).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

p("Bogotá D.C., agosto de 2026", sangria=False).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ---------------------------------------------------------------------------
# Declaración inicial de uso de herramientas de IA
# ---------------------------------------------------------------------------

h1("Declaración de uso de herramientas de inteligencia artificial")

p("En el desarrollo de la Fase 2 de este proyecto, el equipo empleó herramientas de "
  "inteligencia artificial como apoyo en tareas de programación, depuración, elaboración "
  "de documentación técnica y diseño de pruebas. Todo el material generado fue revisado, "
  "probado y comprendido por los integrantes del equipo, quienes asumen la autoría y la "
  "responsabilidad académica de los resultados presentados en este documento.")

doc.add_page_break()

# ---------------------------------------------------------------------------
# Introducción
# ---------------------------------------------------------------------------

h1("Introducción")

p("La guía número 2 de la asignatura Sistemas Distribuidos plantea el desarrollo de un "
  "servicio distribuido de registro y consulta de información compuesto por un cliente y "
  "dos servidores replicados, cada uno con su propia base de datos. En la Fase 1 se elaboró "
  "el diseño de la propuesta: arquitectura, tecnologías y flujos de comunicación. Este "
  "documento corresponde a la Fase 2, en la que se implementó el sistema diseñado y se "
  "recopilaron las evidencias de su funcionamiento.")

p("El sistema está formado por un cliente basado en Ubuntu 22.04, un Servidor 1 sobre "
  "Debian GNU/Linux, un Servidor 2 sobre Fedora Linux y una instancia independiente de "
  "PostgreSQL 16 por servidor (db1 y db2), todo orquestado con Docker Compose sobre una red "
  "interna con direcciones IP fijas. El cliente se comunica con los servidores mediante una "
  "API REST, y los servidores mantienen un canal TCP propio para el heartbeat, la "
  "replicación de registros y la resincronización tras una caída.")

p("El documento presenta los requerimientos del sistema, un resumen de la arquitectura, la "
  "instalación y configuración realizadas, las evidencias de ejecución y de comunicación "
  "entre servidores, los resultados de las pruebas y las conclusiones del trabajo. Las "
  "imágenes están marcadas con etiquetas para insertar las capturas tomadas por el equipo.")

doc.add_page_break()

# ---------------------------------------------------------------------------
# Requerimientos
# ---------------------------------------------------------------------------

h1("Requerimientos")

p("Los requerimientos se derivan de la propuesta de la Fase 1 y de las indicaciones de la "
  "guía de la asignatura.")

h2("Requerimientos funcionales")

for texto in (
    "RF1. Registrar información desde el cliente mediante un menú interactivo o por línea de comandos.",
    "RF2. Enviar cada registro al servidor correspondiente, con conmutación automática al otro servidor si el preferido no responde.",
    "RF3. Consultar la información almacenada desde cualquiera de los dos servidores.",
    "RF4. Mantener comunicación y sincronía constante entre los dos servidores mediante un heartbeat periódico.",
    "RF5. Replicar cada registro creado en un servidor hacia el otro, sin duplicar datos.",
    "RF6. Conservar los registros pendientes de replicar y reintentar su envío mientras el otro servidor no esté disponible.",
    "RF7. Resincronizar un servidor que se recupera de una caída, hasta que ambos vuelvan a tener los mismos datos.",
    "RF8. Identificar en cada respuesta qué servidor atendió la solicitud.",
    "RF9. Mostrar el estado de conexión de cada servidor y su visión del servidor par.",
    "RF10. Registrar en el cliente un historial local en SQLite de las comunicaciones realizadas, sin duplicar los datos de los servidores.",
):
    item(texto)

h2("Requerimientos no funcionales")

for texto in (
    "RNF1. Tolerancia a fallos: el sistema debe seguir atendiendo solicitudes aunque uno de los dos servidores esté caído.",
    "RNF2. Fiabilidad: ninguna escritura aceptada debe perderse; cada registro se guarda primero en la base local y luego se replica.",
    "RNF3. Rendimiento: cada servidor accede a su propia base local y la replicación se intenta de inmediato, con reintentos periódicos.",
    "RNF4. Idempotencia: la replicación debe poder repetirse sin generar duplicados (inserción por UUID con ON CONFLICT DO NOTHING).",
    "RNF5. Portabilidad: el mismo código de servidor debe ejecutarse sin cambios en Debian y en Fedora.",
    "RNF6. Simplicidad y mantenibilidad: código corto, comentado y versionado, con una única implementación compartida por ambos servidores.",
    "RNF7. Aislamiento: la red interna no debe exponer más puertos de los necesarios en el anfitrión.",
):
    item(texto)

doc.add_page_break()

# ---------------------------------------------------------------------------
# Resumen del sistema
# ---------------------------------------------------------------------------

h1("Resumen del sistema")

p("El despliegue consta de cinco contenedores conectados a la red interna sdnet "
  "(172.28.0.0/16): el cliente (Ubuntu 22.04), el Servidor 1 (Debian) con su base db1, y el "
  "Servidor 2 (Fedora) con su base db2. Cada servidor expone una API REST implementada con "
  "Flask y administra su propia instancia de PostgreSQL; el cliente consume ambas API y "
  "conserva un historial local en SQLite. La Figura 1 muestra la arquitectura de despliegue.")

figura(1,
       "Arquitectura de despliegue del sistema distribuido",
       "[DIAGRAMA: docs/FASE-02/diagramas/despliegue.puml (renderizar con PlantUML)]",
       nota="Elaboración propia con PlantUML, a partir del diseño de la Fase 1.")

p("La comunicación cliente-servidor usa REST sobre HTTP. Entre servidores existe además un "
  "canal TCP propio (puertos 6000 y 6001, solo en la red interna) con mensajes JSON de una "
  "línea y tres funciones: un heartbeat cada 5 segundos (PING/PONG) que permite a cada nodo "
  "conocer el estado del par; la replicación por empuje, en la que cada registro local se "
  "envía de inmediato al par (REPLICATE/ACK) y, si el par no responde, queda pendiente en "
  "una tabla outbox y se reintenta cada 5 segundos; y la resincronización por arrastre "
  "(SYNC_REQUEST/SYNC_BATCH), que se ejecuta al arrancar, al detectar la recuperación del "
  "par y cada 60 segundos como red de seguridad. Todas las inserciones son idempotentes por "
  "UUID, de modo que repetir un mensaje nunca duplica datos. La Figura 2 ilustra el flujo "
  "normal de registro y consulta, y la Figura 3 el escenario de caída y recuperación.")

figura(2,
       "Registro y consulta en operación normal",
       "[DIAGRAMA: docs/FASE-02/diagramas/secuencia-registro-consulta.puml (renderizar con PlantUML)]",
       nota="Elaboración propia con PlantUML.")

figura(3,
       "Caída del Servidor 2 y resincronización",
       "[DIAGRAMA: docs/FASE-02/diagramas/secuencia-recuperacion.puml (renderizar con PlantUML)]",
       nota="Elaboración propia con PlantUML.")

doc.add_page_break()

# ---------------------------------------------------------------------------
# Instalación y configuración
# ---------------------------------------------------------------------------

h1("Instalación y configuración")

p("A continuación se resumen los siete puntos solicitados en las instrucciones de la Fase "
  "2. El detalle completo, junto con las variables de configuración y la solución de "
  "problemas, está en el manual de instalación y configuración del repositorio.")

h2("Instalación del lenguaje y del runtime")

p("El lenguaje del proyecto es Python 3, instalado dentro de cada contenedor con el gestor "
  "de paquetes de su distribución: apt en el cliente (Ubuntu 22.04) y en el Servidor 1 "
  "(Debian), y dnf en el Servidor 2 (Fedora). No se instala ningún lenguaje ni runtime en el "
  "anfitrión; la instalación ocurre automáticamente durante la construcción de las imágenes "
  "con Docker Compose.")

h2("Instalación de dependencias")

p("Las dependencias están declaradas en archivos requirements.txt y se instalan con pip "
  "durante el build de cada imagen: Flask y psycopg (driver de PostgreSQL) en los "
  "servidores, y requests en el cliente. En Debian se usa el flag "
  "--break-system-packages porque su Python de sistema está marcado como "
  "externally managed (PEP 668).")

h2("Configuración de red")

p("Docker Compose crea una red bridge interna llamada sdnet con el rango 172.28.0.0/16 y "
  "direcciones IP fijas para cada servicio (Tabla 1). Los servicios se alcanzan entre sí por "
  "sus nombres DNS internos, que son más legibles y estables que las IP.")

titulo_tabla(1, "Direcciones IP y nombres DNS de la red interna sdnet")
tabla(
    ("Servicio", "IP en sdnet", "Nombre DNS interno"),
    (("db1 (PostgreSQL de server1)", "172.28.0.11", "db1"),
     ("server1 (Debian)", "172.28.0.21", "server1"),
     ("db2 (PostgreSQL de server2)", "172.28.0.12", "db2"),
     ("server2 (Fedora)", "172.28.0.22", "server2"),
     ("client (Ubuntu)", "172.28.0.30", "client")),
    anchos=(7.0, 4.5, 4.5),
    nota="Elaboración propia con base en docker-compose.yml del repositorio.",
)

h2("Configuración de puertos")

p("Cada servidor usa dos puertos internos: el de la API REST y el del canal TCP de "
  "sincronía. Solo los puertos de la API se publican en el anfitrión, para probar con el "
  "navegador o con curl; PostgreSQL y el canal TCP permanecen dentro de la red interna "
  "(Tabla 2).")

titulo_tabla(2, "Puertos de cada servicio")
tabla(
    ("Servicio", "Puerto interno", "Publicado en el anfitrión", "Uso"),
    (("server1", "5000 (Flask)", "localhost:5000", "API REST del Servidor 1"),
     ("server1", "6000 (TCP propio)", "—", "Heartbeat, replicación y resincronización"),
     ("server2", "5001 (Flask)", "localhost:5001", "API REST del Servidor 2"),
     ("server2", "6001 (TCP propio)", "—", "Heartbeat, replicación y resincronización"),
     ("db1 / db2", "5432 (PostgreSQL)", "—", "Base local de cada servidor"),
     ("client", "—", "—", "Solo consumidor")),
    anchos=(3.0, 4.5, 4.5, 6.0),
    nota="Elaboración propia con base en docker-compose.yml del repositorio.",
)

h2("Instalación de la base de datos")

p("Cada servidor administra su propia instancia de PostgreSQL 16 mediante la imagen "
  "oficial, con la base records, el usuario sduser y volúmenes nombrados (db1_data y "
  "db2_data) para persistir los datos. El esquema se crea automáticamente al arrancar cada "
  "servidor (sentencias CREATE TABLE IF NOT EXISTS) con dos tablas: records (registros, "
  "clave UUID) y outbox (registros locales pendientes de replicar). El cliente, por su "
  "parte, guarda su historial en una base SQLite dentro de un volumen propio.")

h2("Ejecución del software")

p("El sistema se levanta con un único comando de Docker Compose, seguido de la espera de "
  "los healthchecks; el cliente se ejecuta con un menú interactivo o con subcomandos:")

codigo((
    "docker compose up -d --build db1 db2 server1 server2",
    "docker compose run --rm --no-deps client                         # menú interactivo",
    "docker compose run --rm --no-deps -T client python3 client.py create --payload \"Hola mundo\" --server 1",
    "docker compose run --rm --no-deps -T client python3 client.py list --server 2",
    "docker compose run --rm --no-deps -T client python3 client.py status",
))

h2("Comunicación entre servidores")

p("Los servidores se comunican por un canal TCP propio con mensajes JSON de una línea "
  "(puertos 6000 y 6001). El protocolo tiene cuatro tipos de mensaje: PING/PONG para el "
  "heartbeat cada 5 segundos; REPLICATE/ACK para empujar un registro al par; "
  "SYNC_REQUEST/SYNC_BATCH para pedir los registros del par por páginas de 100, ordenadas "
  "por fecha e identificador; y ERROR para respuestas no reconocidas. La actividad del "
  "canal queda registrada en los logs con el prefijo [sync] y puede observarse con docker "
  "compose logs. La configuración principal de cada servidor se resume en la Tabla 3.")

titulo_tabla(3, "Variables de entorno principales de los servidores")
tabla(
    ("Variable", "server1", "server2", "Significado"),
    (("SERVER_ID", "server1", "server2", "Identidad del nodo (aparece en served_by)"),
     ("PORT", "5000", "5001", "Puerto de la API REST"),
     ("SYNC_PORT", "6000", "6001", "Puerto del canal TCP de sincronía"),
     ("PEER_HOST", "server2", "server1", "Nombre DNS del servidor par"),
     ("HEARTBEAT_SECONDS / REPLICATE_SECONDS", "5 / 5", "5 / 5", "Intervalos de heartbeat y reintentos"),
     ("SYNC_EVERY_SECONDS", "60", "60", "Resincronización periódica de seguridad")),
    anchos=(5.0, 2.5, 2.5, 6.5),
    nota="Elaboración propia con base en docker-compose.yml del repositorio.",
)

doc.add_page_break()

# ---------------------------------------------------------------------------
# Evidencias
# ---------------------------------------------------------------------------

h1("Evidencias")

p("Esta sección reúne los entregables 5, 6 y 7 de las instrucciones de la Fase 2: "
  "evidencias de ejecución en ambos servidores, evidencias de comunicación entre "
  "servidores y resultados de las pruebas. Las capturas fueron tomadas por el equipo "
  "durante la ejecución del sistema con Windows y Docker Desktop; cada recuadro marcado "
  "con [IMAGEN: ...] se sustituye por la captura correspondiente, cuyo nombre de archivo "
  "coincide con la carpeta de evidencias del repositorio.")

h2("Ambiente de prueba")

titulo_tabla(4, "Datos del ambiente de prueba")
tabla(
    ("Campo", "Valor"),
    (("Fecha de ejecución", "(completar por el equipo)"),
     ("Equipo", "(completar por el equipo)"),
     ("Sistema operativo anfitrión", "(completar por el equipo)"),
     ("Docker / Docker Compose", "(completar por el equipo)"),
     ("Estado inicial", "Despliegue limpio antes de comenzar")),
    anchos=(6.0, 12.0),
    nota="Elaboración propia.",
)

h2("Evidencias de ejecución en ambos servidores")

p("La Figura 4 muestra los cuatro contenedores en estado healthy tras el arranque "
  "(docker compose ps). La Figura 5 confirma que cada servidor ejecuta su distribución "
  "correspondiente: Debian GNU/Linux en el Servidor 1 y Fedora Linux en el Servidor 2 "
  "(archivo /etc/os-release). La Figura 6 muestra, además, el funcionamiento directo de "
  "ambas API REST desde el anfitrión con curl.")

figura(4,
       "Estado de los contenedores después del arranque (docker compose ps)",
       "[IMAGEN: ev-01-docker-compose-ps.png — captura de docker compose ps con los cuatro contenedores healthy]",
       nota="Captura de pantalla del equipo. Elaboración propia.")

figura(5,
       "Distribución de cada servidor (/etc/os-release)",
       "[IMAGEN: ev-02-distros.png — captura de /etc/os-release de server1 (Debian) y server2 (Fedora)]",
       nota="Captura de pantalla del equipo. Elaboración propia.")

figura(6,
       "Pruebas directas contra la API REST con curl",
       "[IMAGEN: ev-11-curl-api.png — captura de los endpoints /api/status y /api/records con curl]",
       nota="Captura de pantalla del equipo. Elaboración propia.")

h2("Evidencias de comunicación entre servidores")

p("La Figura 7 muestra la creación de un registro a través del Servidor 1 (la respuesta "
  "indica atendido por server1). La Figura 8 muestra la consulta del mismo registro desde "
  "el Servidor 2 (atendido por server2), lo que evidencia la replicación entre las dos "
  "bases. La Figura 9 muestra las trazas del canal de sincronía en los logs (prefijo "
  "[sync]) y la Figura 10 una comprobación directa del canal TCP con PING/PONG.")

figura(7,
       "Registro creado a través del Servidor 1",
       "[IMAGEN: ev-03-create-server1.png — captura del comando create con --server 1]",
       nota="Captura de pantalla del equipo. Elaboración propia.")

figura(8,
       "Consulta del registro desde el Servidor 2 (replicación entre servidores)",
       "[IMAGEN: ev-04-replicacion-s1-s2.png — captura del comando list con --server 2 mostrando el registro creado en server1]",
       nota="Captura de pantalla del equipo. Elaboración propia.")

figura(9,
       "Trazas del canal de sincronía en los logs (prefijo [sync])",
       "[IMAGEN: ev-12-logs-sync.png — captura de docker compose logs server1 server2 filtrado por [sync]]",
       nota="Captura de pantalla del equipo. Elaboración propia.")

figura(10,
       "Comprobación directa del canal TCP (PING/PONG)",
       "[IMAGEN: ev-13-ping-tcp.png — captura de python3 ping_peer.py desde cada servidor]",
       nota="Captura de pantalla del equipo. Elaboración propia.")

h2("Resultados de las pruebas")

p("Se ejecutaron dos escenarios, automatizados con los scripts de la carpeta tests del "
  "repositorio y verificados paso a paso de forma manual.")

p("Escenario 1, operación normal: se levantó el sistema completo; se creó un registro a "
  "través del Servidor 1; se consultó desde el Servidor 2 (el registro ya estaba replicado); "
  "y se verificó el estado de ambos servidores y su visión del par (Figuras 7, 8 y 11).")

figura(11,
       "Estado de ambos servidores en operación normal",
       "[IMAGEN: ev-05-status-normal.png — captura del comando status con ambos servidores en línea]",
       nota="Captura de pantalla del equipo. Elaboración propia.")

p("Escenario 2, caída y recuperación: se detuvo el Servidor 2 y el heartbeat del Servidor 1 "
  "detectó la caída (Figura 12); se crearon dos registros con el Servidor 2 caído, uno de "
  "ellos pidiendo explícitamente el Servidor 2 para evidenciar la conmutación automática "
  "del cliente (Figura 13); se comprobó que el Servidor 1 seguía atendiendo con normalidad "
  "(Figura 14); se reactivó el Servidor 2 y, tras el heartbeat y la resincronización, este "
  "recuperó los registros creados durante su caída (Figura 15); por último se verificó el "
  "estado final y el historial SQLite del cliente (Figura 16).")

figura(12,
       "Estado con el Servidor 2 detenido (detección por heartbeat)",
       "[IMAGEN: ev-06-server2-caido.png — captura del comando status con el Servidor 2 caído]",
       nota="Captura de pantalla del equipo. Elaboración propia.")

figura(13,
       "Registros creados durante la caída, con conmutación del cliente",
       "[IMAGEN: ev-07-create-con-server2-caido.png — captura de los comandos create con el Servidor 2 caído]",
       nota="Captura de pantalla del equipo. Elaboración propia.")

figura(14,
       "El Servidor 1 atiende con normalidad mientras el par está caído",
       "[IMAGEN: ev-08-server1-normal-con-par-caido.png — captura del comando list en server1 con el par caído]",
       nota="Captura de pantalla del equipo. Elaboración propia.")

figura(15,
       "Servidor 2 resincronizado con los registros creados durante su caída",
       "[IMAGEN: ev-09-resincronizacion-server2.png — captura del comando list en server2 tras la recuperación]",
       nota="Captura de pantalla del equipo. Elaboración propia.")

figura(16,
       "Estado final de los servidores e historial SQLite del cliente",
       "[IMAGEN: ev-10-status-final-historial.png — captura del comando status y del comando history]",
       nota="Captura de pantalla del equipo. Elaboración propia.")

p("La Tabla 5 resume los resultados de las pruebas. La columna de resultado obtenido se "
  "completa con lo observado en las capturas anteriores.")

titulo_tabla(5, "Resultados de las pruebas")
tabla(
    ("Prueba", "Resultado esperado", "Resultado obtenido", "¿Cumple?"),
    (("Arranque de los contenedores", "Cuatro contenedores en estado healthy", "", ""),
     ("Distribuciones de los servidores", "server1: Debian; server2: Fedora", "", ""),
     ("Registro vía Servidor 1", "Respuesta 201 con served_by = server1", "", ""),
     ("Replicación server1 → server2", "El registro aparece al consultar server2", "", ""),
     ("Estado en operación normal", "Ambos en línea y cada uno ve a su par", "", ""),
     ("Detección de la caída", "peer_online = false y el cliente reporta NO RESPONDE", "", ""),
     ("Conmutación del cliente", "Petición a server2 caído atendida por server1", "", ""),
     ("Continuidad del servicio", "server1 atiende registros y consultas con el par caído", "", ""),
     ("Resincronización", "server2 recupera los registros creados durante la caída", "", ""),
     ("Estado final", "Ambos en línea con el mismo número de registros", "", ""),
     ("Historial SQLite del cliente", "Trazas de todas las comunicaciones", "", ""),
     ("Comunicación TCP entre servidores", "Logs [sync] y PONG directo por TCP", "", "")),
    anchos=(4.5, 6.5, 4.5, 2.0),
    nota="Elaboración propia a partir de la ejecución de los escenarios en el equipo.",
)

doc.add_page_break()

# ---------------------------------------------------------------------------
# Conclusiones
# ---------------------------------------------------------------------------

h1("Conclusiones")

p("Se implementó por completo el sistema distribuido diseñado en la Fase 1 y se verificó "
  "su funcionamiento sobre los dos servidores con distribuciones distintas, Debian y "
  "Fedora, ejecutando un mismo código. La replicación activa-activa entre las bases db1 y "
  "db2 permitió que un registro creado en un servidor estuviera disponible de inmediato en "
  "el otro, y que cada respuesta identificara al servidor que la atendió.")

p("La tolerancia a fallos se comprobó en el escenario de caída del Servidor 2: el Servidor "
  "1 detectó la falla por heartbeat y continuó atendiendo al cliente, los registros creados "
  "durante la caída se conservaron como pendientes y, al reactivarse el nodo, el sistema "
  "convergió solo mediante los reintentos del outbox y la resincronización paginada. La "
  "idempotencia por UUID resultó clave: permitió que ambos mecanismos actuaran a la vez sin "
  "duplicar datos.")

p("Del trabajo se rescatan varios aprendizajes: la orquestación con contenedores permite "
  "ejecutar un mismo código en distribuciones distintas y desplegar todo el sistema con un "
  "solo comando; implementar un protocolo de aplicación propio sobre TCP resultó sencillo "
  "y suficiente para el heartbeat, la replicación y la resincronización; y las pruebas de "
  "caída demostraron la importancia de registrar lo pendiente y reintentar en lugar de "
  "fallar. Como límite del diseño, la consistencia es eventual (entre una escritura y su "
  "réplica hay una pequeña ventana de tiempo) y no hay resolución de conflictos, porque "
  "los registros solo se crean y nunca se editan ni eliminan.")

doc.add_page_break()

# ---------------------------------------------------------------------------
# Referencias
# ---------------------------------------------------------------------------

h1("Referencias")

for texto in (
    "Anthropic. (2026). Claude [Modelo de lenguaje de gran tamaño]. https://claude.ai",
    "Coulouris, G., Dollimore, J., Kindberg, T. y Blair, G. (2011). Distributed systems: "
    "Concepts and design (5.ª ed.). Addison-Wesley.",
    "Docker Inc. (s. f.). Docker Compose overview. https://docs.docker.com/compose/",
    "Osorio Tabares, J. J. (2026). Guía #02: Comunicación en sistemas distribuidos "
    "[Material del curso]. Sistemas Distribuidos, Universidad Manuela Beltrán.",
    "Pallets Projects. (s. f.). Flask documentation. https://flask.palletsprojects.com/",
    "PlantUML. (s. f.). PlantUML language reference guide. https://plantuml.com/",
    "PostgreSQL Global Development Group. (s. f.). PostgreSQL 16 documentation. "
    "https://www.postgresql.org/docs/16/",
    "Python Software Foundation. (s. f.). The Python standard library. "
    "https://docs.python.org/3/library/",
    "Tanenbaum, A. S. y Van Steen, M. (2007). Distributed systems: Principles and "
    "paradigms (2.ª ed.). Pearson Prentice Hall.",
):
    referencia(texto)

doc.add_page_break()

# ---------------------------------------------------------------------------
# Declaración final de uso de herramientas de IA
# ---------------------------------------------------------------------------

h1("Declaración final de uso de herramientas de inteligencia artificial")

p("Reiteramos que las herramientas de inteligencia artificial fueron empleadas únicamente "
  "como apoyo durante el desarrollo de la Fase 2, y que el código, las configuraciones, "
  "las pruebas y las evidencias presentadas en este documento fueron ejecutadas y "
  "verificadas por los integrantes del equipo.")

# ---------------------------------------------------------------------------

doc.save(SALIDA)
print(f"Documento generado: {SALIDA}")
